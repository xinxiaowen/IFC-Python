from docx import Document
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Border, Side, Alignment
from openpyxl.utils import get_column_letter

def read_word_tables(word_path):
    """
    读取 Word 文档中的所有表格，返回一个列表。
    每个元素代表一个表格，为包含行数据的二维列表。
    同时记录合并信息、对齐、背景色等（仅做示例，可扩展）。
    """
    doc = Document(word_path)
    
    all_tables = []
    
    for table in doc.tables:
        # 存储当前表格的所有行
        rows_data = []
        
        # 为了演示单元格合并，我们需要获取同一 "Grid" 下的起始行列等信息
        # python-docx 中并没有非常直接的“已合并”标记，需要自己通过底层xml判断
        # 这里示例只演示获取每个cell的text和简单对齐，合并等可在后续扩展
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                grid_span = cell._element.grid_span if cell._element.grid_span is not None else 1
                row_span = cell._element.vMerge if cell._element.vMerge is not None else 1
                cell_info = {
                    "text": cell.text.strip(),
                    "alignment": cell.paragraphs[0].alignment if cell.paragraphs else None,
                    "vertical_alignment": getattr(cell, 'vertical_alignment', None),  # 安全获取 vertical_alignment 属性
                    # 这里示例性地拿到一些底层 xml 信息，比如背景色、边框等，可以继续挖掘
                    "tc_pr": cell._element.tcPr,
                    "grid_span": int(grid_span) if isinstance(grid_span, int) else 1,  # 列合并
                    "row_span": int(row_span) if isinstance(row_span, int) else 1,  # 行合并
                    "merge": cell._element.getchildren()[0].tag.endswith('vMerge')  # 是否合并单元格
                }
                row_cells.append(cell_info)
            rows_data.append(row_cells)
        
        all_tables.append(rows_data)
    # print(all_tables[1][2][3])
    return all_tables


def write_tables_to_excel(tables, excel_path):
    """
    将列表形式的表格数据写入到 Excel，并尝试保留部分格式。
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "WordTables"
    
    current_row = 1  # 在 Excel 中插入表格的起始行
    
    # 为演示简单，将所有表格写到同一个 Sheet 内，并在表格间空一行
    for table_index, table in enumerate(tables, start=1):
        rows_count = len(table)
        if rows_count == 0:
            continue
        
        cols_count = len(table[0])
        
        # 写入表格内容
        for r in range(rows_count):
            for c in range(cols_count):
                cell_info = table[r][c]
                text = cell_info["text"]
                alignment = cell_info["alignment"]
                
                # 写入单元格数据
                excel_cell = ws.cell(row=current_row + r, column=c + 1, value=text)
                
                # 设置对齐（仅举例）
                # Word 中的对齐值：0=LEFT, 1=CENTER, 2=RIGHT, 3=JUSTIFY
                if alignment == 0:
                    excel_cell.alignment = Alignment(horizontal="left")
                elif alignment == 1:
                    excel_cell.alignment = Alignment(horizontal="center")
                elif alignment == 2:
                    excel_cell.alignment = Alignment(horizontal="right")
                elif alignment == 3:
                    excel_cell.alignment = Alignment(justifyLastLine=True)    
                # 其他的可自行加
                
                # 这里可以示例性地解析 cell_info["tc_pr"] 获取更多信息，比如背景色
                # 比如:
                # bg_color = parse_bg_color_from_tcPr(cell_info["tc_pr"])
                # if bg_color:
                #     excel_cell.fill = PatternFill("solid", fgColor=bg_color)
                
                # 设置简单边框（仅示例，可以按需调整）
                thin = Side(border_style="thin", color="000000")
                excel_cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                
                # 字体示例
                excel_cell.font = Font(name="宋体", size=12)
        
        # 处理合并单元格
        merged_cells = set()  # 用于记录已经合并过的单元格

        for r in range(rows_count):
            for c in range(cols_count):
                cell_info = table[r][c]
                grid_span = cell_info["grid_span"]
                row_span = cell_info["row_span"]
                
                if grid_span > 1 or row_span > 1:
                    start_row = current_row + r
                    end_row = start_row + row_span - 1
                    start_col = c + 1
                    end_col = start_col + grid_span - 1
                    
                    # 检查是否已经合并过
                    if (start_row, start_col) not in merged_cells:
                        ws.merge_cells(start_row=start_row, start_column=start_col, end_row=end_row, end_column=end_col)
                    
                    # 记录合并过的单元格
                    for mr in range(start_row, end_row + 1):
                        for mc in range(start_col, end_col + 1):
                            merged_cells.add((mr, mc))
                            
        # 以列的视角处理合并单元格
        # for c in range(cols_count):
        #     for r in range(rows_count):
        #         cell_info = table[r][c]
        #         grid_span = cell_info["grid_span"]
        #         row_span = cell_info["row_span"]
                
        #         if grid_span > 1 or row_span > 1:
        #             start_row = current_row + r
        #             end_row = start_row + row_span - 1
        #             start_col = c + 1
        #             end_col = start_col + grid_span - 1
                    
        #             # 检查是否已经合并过
        #             if (start_row, start_col) not in merged_cells:
        #                 ws.merge_cells(start_row=start_row, start_column=start_col, end_row=end_row, end_column=end_col)
                    
        #             # 记录合并过的单元格
        #             for mr in range(start_row, end_row + 1):
        #                 for mc in range(start_col, end_col + 1):
        #                     merged_cells.add((mr, mc))

        
        # 表格写完后，current_row 往下移动
        current_row += rows_count + 1  # 空 1 行，区隔不同表格
    
    wb.save(excel_path)
    print(f"所有表格已写入 {excel_path} 文件。")


if __name__ == "__main__":
    # 1. 读取Word中的表格
    word_path = r'D:\表B.4.30.docx'
    tables_data = read_word_tables(word_path)
    
    # 2. 将表格数据写入到Excel
    excel_path = r'D:\表B.4.30.xlsx'
    write_tables_to_excel(tables_data, excel_path)
